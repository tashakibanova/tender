"""Document preparation from templates."""

from __future__ import annotations

from pathlib import Path
from typing import Dict

from docx import Document

from .storage import LocalStorage
from utils.file_parser import open_file_for_edit


class DocumentPreparationService:
    def __init__(self, storage: LocalStorage) -> None:
        self.storage = storage

    def fill_template(self, template_path: str, output_path: str, data: Dict[str, str]) -> str:
        document = Document(template_path)
        self._replace_placeholders(document, data)
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_file)
        open_file_for_edit(str(output_file))
        return str(output_file)

    def _replace_placeholders(self, document: Document, data: Dict[str, str]) -> None:
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, data)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data)

    def _replace_in_paragraph(self, paragraph, data: Dict[str, str]) -> None:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
